"""
Nexon Workspace Engine for Study Mode / Cross-Document Analysis
===============================================================

Provides lightweight, memory-conscious workspace management for Termux / mobile devices.
Handles quota enforcement, zip extraction (filtering APKs/binaries), structured text/table
extraction (ignoring images), and disk-backed chunk retrieval.
"""

import os
import shutil
import zipfile
import json
import re
from pathlib import Path
from typing import List, Dict, Any, Optional

# Workspace configuration and mobile limits
WORKSPACE_DIR = os.path.expanduser(os.getenv("NEXON_WORKSPACE_DIR", "~/nexon_workspace"))
MAX_WORKSPACE_MB = 150    # Total workspace disk limit (MB)
MAX_FILE_SIZE_MB = 150    # Max single upload file size (MB)
CHUNK_SIZE_CHARS = 1200   # ~250-300 words per chunk
CHUNK_OVERLAP = 150       # Overlap between consecutive chunks

# Prohibited executable/binary/media/image formats
BLOCKED_EXTENSIONS = {
    ".apk", ".exe", ".dll", ".so", ".bin", ".iso", ".img",
    ".deb", ".rpm", ".dmg", ".class", ".pyc", ".o", ".a", ".elf",
    ".mp3", ".mp4", ".wav", ".avi", ".mov", ".flv", ".wmv", ".webm",
    ".ogg", ".flac", ".aac", ".m4a", ".m4v", ".mkv",
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp",
    ".svg", ".ico", ".psd", ".ai", ".raw"
}

ALLOWED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".json",
    ".rtf", ".odt", ".pptx", ".ppt", ".xlsx", ".xls",
    ".epub", ".html", ".htm", ".xml", ".log", ".tex"
}


class WorkspaceManager:
    """Memory-efficient document workspace engine for Termux/mobile environments."""

    def __init__(self, workspace_path: str = WORKSPACE_DIR):
        self.workspace_path = Path(workspace_path)
        self.workspace_path.mkdir(parents=True, exist_ok=True)
        self.index_file = self.workspace_path / ".workspace_index.json"

    def check_storage_quota(self, incoming_bytes: int = 0) -> Dict[str, Any]:
        """Ensures mobile device storage and workspace quota limits are respected."""
        total_used = sum(
            f.stat().st_size for f in self.workspace_path.rglob('*')
            if f.is_file() and not f.is_symlink()
        )
        total_used_mb = total_used / (1024 * 1024)
        incoming_mb = incoming_bytes / (1024 * 1024)

        if (total_used_mb + incoming_mb) > MAX_WORKSPACE_MB:
            return {
                "allowed": False,
                "reason": f"Workspace limit exceeded ({total_used_mb:.1f}MB used out of {MAX_WORKSPACE_MB}MB quota)."
            }

        usage = shutil.disk_usage(self.workspace_path)
        free_mb = usage.free / (1024 * 1024)
        if free_mb < 200:
            return {
                "allowed": False,
                "reason": f"Low device storage ({free_mb:.1f}MB free). Clear device storage to continue."
            }

        return {
            "allowed": True,
            "used_mb": round(total_used_mb, 2),
            "quota_mb": MAX_WORKSPACE_MB,
            "free_disk_mb": round(free_mb, 2)
        }

    def ingest_file(self, file_path: str, rebuild: bool = True) -> Dict[str, Any]:
        """Ingests a file or unzips archives safely, discarding prohibited binaries."""
        path = Path(file_path)
        if not path.exists():
            return {"status": "error", "message": f"File not found: {file_path}"}

        file_size_mb = path.stat().st_size / (1024 * 1024)
        if file_size_mb > MAX_FILE_SIZE_MB:
            return {"status": "error", "message": f"File exceeds maximum single file limit of {MAX_FILE_SIZE_MB}MB"}

        quota_check = self.check_storage_quota(path.stat().st_size)
        if not quota_check["allowed"]:
            return {"status": "error", "message": quota_check["reason"]}

        extracted_files = []

        # Unpack Zip
        if path.suffix.lower() == ".zip":
            try:
                with zipfile.ZipFile(path, 'r') as zip_ref:
                    for member in zip_ref.infolist():
                        member_path = Path(member.filename)

                        # Prevent Zip-Slip
                        target_path = (self.workspace_path / member_path).resolve()
                        if not str(target_path).startswith(str(self.workspace_path.resolve())):
                            continue

                        # Reject APKs and binary formats
                        if member_path.suffix.lower() in BLOCKED_EXTENSIONS:
                            continue

                        if not member.is_dir() and member_path.suffix.lower() in ALLOWED_EXTENSIONS:
                            zip_ref.extract(member, self.workspace_path)
                            extracted_files.append(str(member_path))

                # Remove original zip archive to save mobile disk space
                if path.exists() and str(path.parent.resolve()) == str(self.workspace_path.resolve()):
                    os.remove(path)

                index_res = self.rebuild_index() if rebuild else {"total_files": 0, "total_chunks": 0}
                return {
                    "status": "success",
                    "action": "unzipped",
                    "extracted_files": extracted_files,
                    "index_summary": index_res
                }
            except Exception as e:
                return {"status": "error", "message": f"Failed to unpack zip archive: {str(e)}"}
        else:
            if path.suffix.lower() in BLOCKED_EXTENSIONS:
                return {"status": "error", "message": f"Security restriction: {path.suffix} binary files are blocked."}

            if path.suffix.lower() not in ALLOWED_EXTENSIONS:
                return {"status": "error", "message": f"Unsupported format {path.suffix}. Allowed: {ALLOWED_EXTENSIONS}"}

            dest = self.workspace_path / path.name
            if path.resolve() != dest.resolve():
                shutil.copy(path, dest)

            index_res = self.rebuild_index() if rebuild else {"total_files": 0, "total_chunks": 0}
            return {
                "status": "success",
                "action": "ingested",
                "file": path.name,
                "index_summary": index_res
            }

    def extract_text_content(self, file_path: Path) -> str:
        """Extracts readable text and tables while ignoring bitmap images."""
        ext = file_path.suffix.lower()

        if ext in [".txt", ".md", ".csv", ".json"]:
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            except Exception as e:
                return f"[Error reading text file: {e}]"

        elif ext == ".pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                pages = []
                for idx, page in enumerate(reader.pages):
                    # Extract text only — images are explicitly skipped
                    page_text = page.extract_text() or ""
                    pages.append(f"\n--- [Page {idx + 1}] ---\n{page_text}")
                return "\n".join(pages)
            except ImportError:
                if shutil.which("pdftotext"):
                    # pdftotext extracts text only, images are ignored
                    out_txt = file_path.with_suffix(".tmp.txt")
                    os.system(f"pdftotext '{file_path}' '{out_txt}'")
                    if out_txt.exists():
                        res = out_txt.read_text(errors="ignore")
                        out_txt.unlink()
                        return res
                return "[PDF Extractor unavailable: Install pypdf or poppler in Termux]"

        elif ext == ".docx":
            try:
                import docx
                doc = docx.Document(file_path)
                content = []

                # Paragraph text
                for p in doc.paragraphs:
                    if p.text.strip():
                        content.append(p.text)

                # Preserve tables in Markdown format
                for t_idx, table in enumerate(doc.tables):
                    content.append(f"\n[Table {t_idx + 1}]")
                    for row in table.rows:
                        row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        content.append("| " + " | ".join(row_cells) + " |")
                    content.append("")

                return "\n".join(content)
            except ImportError:
                return "[docx Extractor unavailable: Install python-docx in Termux]"

        return ""

    def _chunk_text(self, text: str, file_name: str, rel_path: str) -> List[Dict[str, Any]]:
        """Structure-aware chunking: splits by paragraphs/headers, merges small ones."""
        chunks = []
        lines = text.split('\n')
        current_chunk = []
        current_len = 0
        
        def flush_chunk():
            nonlocal current_chunk, current_len
            if current_chunk:
                chunk_text = '\n'.join(current_chunk).strip()
                if chunk_text:
                    chunks.append({
                        "file": file_name,
                        "relative_path": rel_path,
                        "chunk_id": len(chunks),
                        "content": chunk_text
                    })
                current_chunk = []
                current_len = 0

        for line in lines:
            stripped = line.strip()
            # Structural boundaries
            is_header = stripped.startswith('#') or stripped.startswith('--- [Page')
            is_table_row = stripped.startswith('|')
            
            # If we hit a structural boundary and current chunk is big enough, flush
            if (is_header or is_table_row) and current_len > CHUNK_SIZE_CHARS * 0.5:
                flush_chunk()
            
            current_chunk.append(line)
            current_len += len(line) + 1  # +1 for newline
            
            # If chunk exceeds max size, flush
            if current_len >= CHUNK_SIZE_CHARS:
                flush_chunk()
                
        flush_chunk() # Flush any remaining
        
        # Post-processing: if any chunk is STILL too big (e.g., a single massive line),
        # split it manually by characters.
        final_chunks = []
        for c in chunks:
            if len(c["content"]) > CHUNK_SIZE_CHARS * 2:
                text_content = c["content"]
                step = CHUNK_SIZE_CHARS - CHUNK_OVERLAP
                for i in range(0, len(text_content), step):
                    chunk_text = text_content[i : i + CHUNK_SIZE_CHARS]
                    final_chunks.append({
                        "file": file_name,
                        "relative_path": rel_path,
                        "chunk_id": len(final_chunks),
                        "content": chunk_text
                    })
            else:
                c["chunk_id"] = len(final_chunks) # Reindex
                final_chunks.append(c)
                
        return final_chunks

    def check_dependencies(self) -> Dict[str, Any]:
        """Checks for required system and python dependencies for text extraction."""
        deps = {
            "pypdf": False,
            "python-docx": False,
            "pdftotext": False
        }
        try:
            import pypdf
            deps["pypdf"] = True
        except ImportError:
            pass
        try:
            import docx
            deps["python-docx"] = True
        except ImportError:
            pass
        
        if shutil.which("pdftotext"):
            deps["pdftotext"] = True
            
        missing = [k for k, v in deps.items() if not v]
        commands = []
        if "pypdf" in missing:
            commands.append("pip install pypdf")
        if "python-docx" in missing:
            commands.append("pip install python-docx")
        if "pdftotext" in missing:
            commands.append("pkg install poppler")
            
        return {
            "dependencies": deps,
            "missing": missing,
            "commands": commands,
            "all_present": len(missing) == 0
        }

    def rebuild_index(self) -> Dict[str, Any]:
        """Chunk workspace documents and save chunk index to disk."""
        chunks = []
        failed_files = []
        files = [f for f in self.workspace_path.rglob('*') if f.is_file() and not f.name.startswith('.')]

        for f in files:
            if f.suffix.lower() in ALLOWED_EXTENSIONS:
                try:
                    text = self.extract_text_content(f)
                    if not text.strip():
                        continue
                    chunks.extend(self._chunk_text(text, f.name, str(f.relative_to(self.workspace_path))))
                except Exception as exc:
                    # One malformed document must not kill the whole reindex;
                    # record it and continue so the app still gets a success signal.
                    failed_files.append({"file": f.name, "reason": str(exc)})

        with open(self.index_file, "w", encoding="utf-8") as idx_f:
            json.dump(chunks, idx_f)

        return {"total_files": len(files), "total_chunks": len(chunks), "failed_files": failed_files}

    def read_page(self, file_path: str, page: int = 1) -> Dict[str, Any]:
        """Read a specific page of a document (text only, images skipped)."""
        path = Path(file_path)
        if not path.exists():
            # Try relative to workspace
            path = self.workspace_path / file_path
        if not path.exists():
            return {"status": "error", "message": f"File not found: {file_path}"}

        ext = path.suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            return {"status": "error", "message": f"Unsupported format {ext}"}

        if ext == ".pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(str(path))
                total_pages = len(reader.pages)
                if page < 1 or page > total_pages:
                    return {"status": "error", "message": f"Page {page} out of range (1-{total_pages})"}
                page_text = reader.pages[page - 1].extract_text() or ""
                return {
                    "status": "success",
                    "file": path.name,
                    "page": page,
                    "total_pages": total_pages,
                    "content": page_text.strip()
                }
            except ImportError:
                if shutil.which("pdftotext"):
                    import subprocess
                    result = subprocess.run(
                        ["pdftotext", "-f", str(page), "-l", str(page), str(path), "-"],
                        capture_output=True, text=True
                    )
                    return {
                        "status": "success",
                        "file": path.name,
                        "page": page,
                        "content": result.stdout.strip()
                    }
                return {"status": "error", "message": "PDF extractor unavailable"}
        elif ext in [".docx", ".doc"]:
            try:
                import docx
                doc = docx.Document(str(path))
                # Approximate pages by paragraph count (no native page concept)
                content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                return {
                    "status": "success",
                    "file": path.name,
                    "page": page,
                    "total_pages": 1,
                    "content": content
                }
            except ImportError:
                return {"status": "error", "message": "python-docx unavailable"}
        else:
            # For text files, return entire content
            text = self.extract_text_content(path)
            return {
                "status": "success",
                "file": path.name,
                "page": 1,
                "total_pages": 1,
                "content": text.strip()
            }

    def get_outline(self, file_path: str) -> Dict[str, Any]:
        """Get document structure (headings, chapters, outline) — text only."""
        path = Path(file_path)
        if not path.exists():
            path = self.workspace_path / file_path
        if not path.exists():
            return {"status": "error", "message": f"File not found: {file_path}"}

        ext = path.suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            return {"status": "error", "message": f"Unsupported format {ext}"}

        if ext == ".pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(str(path))
                outline = []
                total_pages = len(reader.pages)

                def _walk_outline(items, depth=0):
                    for item in items:
                        if isinstance(item, list):
                            _walk_outline(item, depth + 1)
                        else:
                            title = item.title if hasattr(item, 'title') else str(item)
                            page_num = None
                            try:
                                page_num = reader.get_destination_page_number(item) + 1
                            except Exception:
                                pass
                            outline.append({
                                "title": title,
                                "page": page_num,
                                "depth": depth
                            })

                if reader.outline:
                    _walk_outline(reader.outline)

                if not outline:
                    # Fallback: generate outline from text content (first lines of each page)
                    for idx, pg in enumerate(reader.pages):
                        text = (pg.extract_text() or "").strip()[:200]
                        if text:
                            outline.append({
                                "title": text.split("\n")[0][:100],
                                "page": idx + 1,
                                "depth": 0
                            })

                return {
                    "status": "success",
                    "file": path.name,
                    "total_pages": total_pages,
                    "outline": outline
                }
            except ImportError:
                return {"status": "error", "message": "pypdf unavailable for outline"}
        elif ext == ".docx":
            try:
                import docx
                doc = docx.Document(str(path))
                outline = []
                for p in doc.paragraphs:
                    if p.style and p.style.name and "Heading" in p.style.name:
                        try:
                            level = int(p.style.name.split()[-1])
                        except (ValueError, IndexError):
                            level = 1
                        outline.append({"title": p.text, "page": None, "depth": level - 1})
                return {
                    "status": "success",
                    "file": path.name,
                    "total_pages": 1,
                    "outline": outline
                }
            except ImportError:
                return {"status": "error", "message": "python-docx unavailable"}
        else:
            # For text/md files, extract headings
            text = self.extract_text_content(path)
            lines = text.split("\n")
            outline = []
            for idx, line in enumerate(lines):
                stripped = line.strip()
                if stripped.startswith("#"):
                    level = len(stripped) - len(stripped.lstrip("#"))
                    outline.append({
                        "title": stripped.lstrip("# ").strip(),
                        "page": None,
                        "depth": level - 1
                    })
                elif stripped and len(stripped) < 100 and idx == 0:
                    outline.append({
                        "title": stripped,
                        "page": None,
                        "depth": 0
                    })
            return {
                "status": "success",
                "file": path.name,
                "total_pages": 1,
                "outline": outline
            }

    def search_chunks(self, query, top_k: int = 5) -> List[Dict[str, Any]]:
        """Search indexed chunks from disk using keyword scoring."""
        if not self.index_file.exists():
            self.rebuild_index()

        try:
            with open(self.index_file, "r", encoding="utf-8") as idx_f:
                chunks = json.load(idx_f)
        except Exception:
            return []

        # Support array of queries for batch RAG
        if isinstance(query, list):
            query = " ".join(str(q) for q in query)

        keywords = [k.lower() for k in re.findall(r'\w+', query) if len(k) > 2]
        if not keywords:
            return chunks[:top_k]

        scored = []
        for chunk in chunks:
            text_lower = chunk["content"].lower()
            score = sum(text_lower.count(kw) for kw in keywords)
            if score > 0:
                scored.append((score, chunk))

        scored.sort(key=lambda x: x[0], reverse=True)
        return [item[1] for item in scored[:top_k]]

    def list_files(self) -> Dict[str, Any]:
        """Returns workspace file inventory and quota details."""
        files = []
        total_size = 0
        for f in self.workspace_path.rglob('*'):
            if f.is_file() and not f.name.startswith('.'):
                sz = f.stat().st_size
                total_size += sz
                files.append({
                    "name": f.name,
                    "path": str(f.relative_to(self.workspace_path)),
                    "size_kb": round(sz / 1024, 1),
                    "ext": f.suffix.lower()
                })

        quota = self.check_storage_quota()
        return {
            "workspace_dir": str(self.workspace_path),
            "file_count": len(files),
            "total_used_mb": round(total_size / (1024 * 1024), 2),
            "quota_mb": MAX_WORKSPACE_MB,
            "files": files
        }


if __name__ == "__main__":
    import sys
    mgr = WorkspaceManager()
    if len(sys.argv) > 1:
        cmd = sys.argv[1]
        if cmd == "list":
            print(json.dumps(mgr.list_files(), indent=2))
        elif cmd == "ingest" and len(sys.argv) > 2:
            print(json.dumps(mgr.ingest_file(sys.argv[2]), indent=2))
        elif cmd == "search" and len(sys.argv) > 2:
            print(json.dumps(mgr.search_chunks(" ".join(sys.argv[2:])), indent=2))
        else:
            print(json.dumps({"error": f"Unknown command {cmd}"}))
    else:
        print(json.dumps(mgr.list_files(), indent=2))
